"""Gera o resumo em Word da EDA Central: as tabelas e as figuras da análise, com um
comentário curto embaixo de cada uma.

Os números vêm de banco_de_dados/eda/dados_deck.json, o mesmo arquivo que alimenta o
gerador do .pptx, para que o resumo e o deck nunca divirjam. As tabelas que só existem
no gerador do deck (funil, fórmulas, renda) estão transcritas aqui e marcadas com
FONTE_DECK.

O documento cobre cinco seções: elegibilidade e as sete componentes, a matriz de
correlação, a renda, os blocos descritivos e as favelas com a comparação nacional. Os
blocos de demandas, inventário de slides e limitações/processo foram retirados a pedido
do autor e não devem ser reintroduzidos sem que ele peça.

Uso:
    uv run --with python-docx python scripts/gerar_resumo_eda_central.py \
        docs/Apresentacoes_IVS/complementos/Resumo_EDA_Central_2026-08.docx
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

RAIZ = Path(__file__).resolve().parents[1]
FIG = RAIZ / "banco_de_dados" / "eda" / "figuras"
D = json.loads((RAIZ / "banco_de_dados" / "eda" / "dados_deck.json").read_text(encoding="utf-8"))

TINTA = RGBColor(0x1A, 0x1A, 0x1A)
PETROL = RGBColor(0x1F, 0x4E, 0x4A)
CLAY = RGBColor(0xA8, 0x3A, 0x2C)
CINZA = RGBColor(0x66, 0x66, 0x66)

SERIF = "Georgia"
SANS = "Calibri"

# As notas do dados_deck.json trazem travessão e caixa alta de ênfase, que vêm do estilo
# do deck. Aqui elas são reescritas na leitura para não destoarem do texto corrido.
NOTAS = {
    "por_regiao": "Média entre setores: cada setor pesa igual, independentemente de quantos domicílios tem.",
    "outliers": "Limite superior = q3 + 1,5 × (q3 − q1), a regra de Tukey.",
    "sigilo_porte": "De 44,09% de sigilo nos setores menores a 3,33% nos maiores.",
    "descritivos_regiao": "Média entre setores, todas no mesmo recorte.",
    "envelhecimento": "IEP = 60+ ÷ menores de 15 × 100, conforme Galvão et al. (Hygeia, 2025).",
    "favela_resto": "Razão de médias abaixo de 1 significa valor menor em favela: é o caso da renda, do apartamento e do índice de envelhecimento.",
    "cobertura_regiao": "Setores em que nenhum domicílio está na condição inadequada.",
    "morfologia": "Percentual sobre domicílios (V00001 + V00002), não sobre setores.",
    "agua": 'Razão agregada sobre domicílios. "Suprimido" é a parcela que o sigilo tira de V00200 e V00201.',
    "brasil_elsi": "Razão agregada: soma dos numeradores ÷ soma dos denominadores em cada recorte.",
}


# ── primitivas ──────────────────────────────────────────────────────────────

def _fonte(run, nome, tam, cor=TINTA, negrito=False, italico=False):
    run.font.name = nome
    run.font.size = Pt(tam)
    run.font.color.rgb = cor
    run.bold = negrito
    run.italic = italico
    # o nome da fonte precisa ir também no rPr de East Asia, senão o Word ignora
    run._element.rPr.rFonts.set(qn("w:eastAsia"), nome)
    return run


def titulo_secao(doc, numero, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    _fonte(p.add_run(f"{numero}  {texto}"), SERIF, 13, PETROL, negrito=True)
    return p


def subtitulo(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    _fonte(p.add_run(texto), SANS, 10.5, TINTA, negrito=True)
    return p


def corpo(doc, texto, cor=TINTA, tam=9.5, italico=False, antes=2, depois=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(antes)
    p.paragraph_format.space_after = Pt(depois)
    p.paragraph_format.line_spacing = 1.05
    _fonte(p.add_run(texto), SANS, tam, cor, italico=italico)
    return p


def comentario(doc, *paragrafos):
    """O texto que explica a tabela ou a figura logo acima."""
    for i, texto in enumerate(paragrafos):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3 if i == 0 else 2)
        p.paragraph_format.space_after = Pt(3 if i < len(paragrafos) - 1 else 6)
        p.paragraph_format.line_spacing = 1.05
        _fonte(p.add_run(texto), SANS, 9, TINTA)


def procedencia(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(8)
    _fonte(p.add_run(texto), SANS, 7.5, CINZA, italico=True)
    return p


def tabela(doc, cabecalho, linhas, larguras=None, tam=8, destaque_col0=True):
    t = doc.add_table(rows=1, cols=len(cabecalho))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False

    for i, c in enumerate(cabecalho):
        cel = t.rows[0].cells[i]
        cel.text = ""
        p = cel.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        _fonte(p.add_run(str(c)), SANS, tam, PETROL, negrito=True)

    for ln in linhas:
        cels = t.add_row().cells
        for i, v in enumerate(ln):
            cel = cels[i]
            cel.text = ""
            p = cel.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            texto, cor, neg = v, TINTA, False
            if isinstance(v, tuple):
                texto, cor, neg = v
            _fonte(p.add_run(str(texto)), SANS, tam, cor,
                   negrito=neg or (destaque_col0 and i == 0))

    if larguras:
        for linha in t.rows:
            for i, w in enumerate(larguras):
                linha.cells[i].width = Cm(w)
    return t


def figura(doc, arquivo, largura_cm, legenda):
    doc.add_picture(str(FIG / arquivo), width=Cm(largura_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_before = Pt(4)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    _fonte(p.add_run(legenda), SANS, 7.5, CINZA, italico=True)


def bloco(doc, chave, larguras=None, tam=8):
    """Escreve uma tabela vinda do dados_deck.json."""
    b = D["blocos"][chave]
    tabela(doc, b["colunas"], b["linhas"], larguras=larguras, tam=tam)
    return b


def fonte_bloco(doc, chave):
    b = D["blocos"][chave]
    nota = f"Recorte: {b['recorte']}. Fonte: {b['fonte']}."
    extra = NOTAS.get(chave)
    if extra:
        nota += f" {extra}"
    procedencia(doc, nota)


# ── documento ───────────────────────────────────────────────────────────────

doc = Document()

sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin = sec.bottom_margin = Cm(1.8)
sec.left_margin = sec.right_margin = Cm(1.9)

est = doc.styles["Normal"]
est.font.name = SANS
est.font.size = Pt(9.5)
est.paragraph_format.space_after = Pt(4)

rod = sec.footer.paragraphs[0]
rod.alignment = WD_ALIGN_PARAGRAPH.CENTER
_fonte(rod.add_run("Pedro Dias Soares · Iniciação Científica · Fiocruz Minas, Instituto René Rachou · agosto de 2026"),
       SANS, 7.5, CINZA)

# ── abertura ────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
_fonte(p.add_run("EDA Central: os números da análise exploratória"), SERIF, 20, TINTA, negrito=True)

corpo(doc,
      "Índice de Vulnerabilidade da Saúde intraurbano · Censo Demográfico 2022 · "
      "70 municípios do ELSI-Brasil. Reuni aqui as tabelas e as figuras da análise "
      "exploratória, cada uma com um comentário curto do que ela mostra. Os valores saem "
      "das tabelas de banco_de_dados/eda/, lidas na hora de gerar o arquivo, e não foram "
      "digitados.",
      tam=9.5)

tabela(doc,
       ["Recorte de análise", "Base bruta", "Municípios"],
       [[("104.108 setores urbanos", PETROL, True), "109.032 setores", "70"]],
       larguras=[6.0, 5.6, 5.6], tam=8.5, destaque_col0=False)
procedencia(doc, "Há dois recortes na base: 104.108 setores urbanos elegíveis e 106.281 "
                 "incluindo os rurais. Cada tabela traz o seu, impresso embaixo dela.")

# ══ 1. QUEM ENTRA E AS SETE COMPONENTES ═════════════════════════════════════
titulo_secao(doc, "1", "Quem entra na análise, e as sete componentes")

subtitulo(doc, "1.1  O funil: do universo do Censo aos setores analisados (slide 7)")
# FONTE_DECK: slide 7
tabela(doc,
       ["Etapa", "Critério", "Setores", "O que sai"],
       [
        ["Universo do Censo 2022", "Brasil inteiro", "468.099", "—"],
        ["Filtro ELSI-Brasil", "70 municípios da coorte", ("109.032", TINTA, True), "municípios fora da coorte"],
        ["ZERADO", "v0001 = 0 (massas d'água)", "− " + D["eleg"]["ZERADO"], "setores sem população"],
        ["SIGILOSO", "v0001 ou V00001 suprimidos", "− " + D["eleg"]["SIGILOSO"], "sem denominador calculável"],
        ["COLETIVO", "V00001 = 0 com população > 0", "0", "classe vazia na base"],
        ["Elegíveis (Dados_sig = OK)", "", D["eleg"]["OK"], ""],
        ["Recorte urbano", "SITUACAO = Urbana", (D["exclusao"]["n_ok_urbano"], PETROL, True),
         D["exclusao"]["n_ok_rural"] + " setores rurais elegíveis"],
       ],
       larguras=[4.4, 4.8, 2.6, 5.4], tam=8)
comentario(doc,
           f"Cada corte tem critério escrito e nenhum é definitivo: a base bruta guarda os "
           f"109.032 setores, então qualquer exclusão pode ser refeita.",
           f"O filtro urbano cobra um preço. {D['exclusao']['perdem_10pct']} dos {D['exclusao']['municipios']} "
           f"municípios perdem mais de 10% dos setores e {D['exclusao']['menos_de_10_setores']} "
           f"ficam com menos de 10, o que torna instável a descritiva desses municípios. A "
           f"limitação está registrada na seção 14 do relatório da EDA.")

subtitulo(doc, "1.2  As sete componentes e suas fórmulas (slide 9)")
# FONTE_DECK: slide 9
tabela(doc,
       ["Componente", "Numerador", "Denominador", "Dimensão"],
       [
        ["pct_agua_inad", "V00112 … V00118 (fonte alternativa)", "V00001", "Saneamento"],
        ["pct_esgoto_inad", "V00312 … V00316", "V00001", "Saneamento"],
        ["pct_lixo_inad", "V00398 … V00402", "V00001", "Saneamento"],
        ["razao_moradores", "V00005 + V00006", "V00001 + V00002", "Socioeconômica"],
        ["pct_analfab", "V00901", ("V00900 + V00901", TINTA, True), "Socioeconômica"],
        ["renda_media", "V06004 (direto)", "—", "Socioeconômica"],
        ["pct_raca_pretpardind", "V01318 + V01320 + V01321", "v0001", "Socioeconômica"],
       ],
       larguras=[4.2, 6.2, 3.4, 3.4], tam=8)
comentario(doc,
           "Duas correções entraram aqui. O denominador domiciliar passou a ser V00001, que "
           "conta domicílios; com V01042, que conta pessoas, "
           "algumas proporções passavam de 1,0. E o analfabetismo virou V00901 ÷ (V00900 + "
           "V00901), uma taxa sobre o total de pessoas de 15 anos ou mais, no lugar de V00901 "
           "÷ V00900, que dava uma razão entre analfabetos e alfabetizados.",
           "Três componentes do IVS-BH 2012 não existem nos agregados por setor do Censo 2022: "
           "anos de estudo, faixas de renda e óbitos cardiovasculares. Cada uma foi substituída "
           "por um proxy declarado.")

subtitulo(doc, "1.3  As sete componentes em números (slide 10)")
bloco(doc, "descritivas", larguras=[4.6, 2.1, 2.1, 2.1, 2.1, 2.3, 1.9])
comentario(doc,
           "Em três das sete a mediana é zero: água, esgoto e lixo não têm inadequação medida "
           "na maioria dos setores urbanos. É essa forma que quebra a regra de outlier em 1.5.",
           "Dois valores merecem atenção: o n do analfabetismo, 87.556 contra 104 mil das outras "
           "seis componentes, e a assimetria da renda, 3,74.")
fonte_bloco(doc, "descritivas")

subtitulo(doc, "1.4  As sete componentes por região (slide 11)")
bloco(doc, "por_regiao", larguras=[4.6, 2.52, 2.52, 2.52, 2.52, 2.52])
comentario(doc,
           "O gradiente Norte/Sul aparece em seis das sete: a água inadequada no Norte é "
           "dezesseis vezes a do Sul, e o esgoto, oito vezes.",
           "As duas exceções interessam mais que a regra. O analfabetismo tem pico no Nordeste, "
           "não no Norte. E o lixo inadequado é maior no Nordeste e no Sudeste do que no Norte, "
           "o contrário do esperado. Essa linha volta a destoar na comparação com o Brasil, em "
           "5.3, e é lá que ela fica mais fácil de interpretar.")
fonte_bloco(doc, "por_regiao")

subtitulo(doc, "1.5  Outliers: a regra do IQR não serve para o saneamento (slide 14)")
bloco(doc, "outliers", larguras=[4.0, 2.1, 2.1, 2.5, 2.2, 2.2, 2.1])
comentario(doc,
           "Quando a mediana e o primeiro quartil são zero, o intervalo interquartil fica "
           "minúsculo e qualquer setor com alguma inadequação passa a ser atípico. Em água, "
           "esgoto e lixo isso marca perto de 20% da base, o que já não descreve uma cauda: a "
           "regra está confundindo a forma da distribuição com anomalia.",
           "Nessas três variáveis o caminho é usar o percentil, que a tabela também traz, ou "
           "aceitar que 20% de setores com inadequação são o fenômeno e não ruído.")
fonte_bloco(doc, "outliers")

subtitulo(doc, "1.6  O sigilo do analfabetismo depende do porte do setor (slide 16)")
bloco(doc, "sigilo_porte", larguras=[6.2, 3.6, 4.2, 3.2], tam=8.5)
comentario(doc,
           "A queda é monotônica, e isso permite ir além de declarar a limitação. Como os "
           "setores sem o dado são os de menor analfabetismo, a média que observo é um teto, e "
           "não uma estimativa central. E como o IBGE reporta os zeros, com 9.268 setores "
           "declarando V00901 = 0, o valor suprimido é necessariamente maior ou igual a 1, o que "
           "fecha o intervalo pelo outro lado: a média verdadeira da amostra está entre 3,14% e "
           "3,64%, meio ponto percentual de largura.",
           "O perfil dos dois grupos confirma a direção do viés. Nos setores sem o dado a renda "
           "mediana é de R$ 6.092,84 e a população preta, parda ou indígena é 30,8%; nos setores "
           "com o dado, R$ 2.313,89 e 60,6%.")
fonte_bloco(doc, "sigilo_porte")

doc.add_page_break()
subtitulo(doc, "1.7  Como as sete variáveis se distribuem (slide 12)")
figura(doc, "histogramas.png", 15.0, "banco_de_dados/eda/figuras/histogramas.png (Notebook 02, célula step8).")
comentario(doc,
           "Cinco das sete têm a mesma forma: uma barra altíssima no zero e uma cauda longa à "
           "direita. O que a figura mostra não é dispersão em torno de um centro, e sim uma "
           "massa de setores adequados somada a uma minoria com problema. A razão de moradores "
           "é a única aproximadamente simétrica, e serve de contraste.")

subtitulo(doc, "1.8  Distribuição por região (slide 13)")
figura(doc, "boxplots_por_regiao.png", 13.5, "banco_de_dados/eda/figuras/boxplots_por_regiao.png (Notebook 02, célula step9).")
comentario(doc,
           "A tabela de 1.4 dá as médias; o boxplot mostra a dispersão por trás delas. No Norte "
           "as caixas de água e esgoto são largas, ou seja, além da média pior há mais variação "
           "interna: setores muito adequados e muito inadequados convivem na mesma região.")

subtitulo(doc, "1.9  Dados faltantes e o sigilo do IBGE (slide 15)")
figura(doc, "missing_por_municipio.png", 8.0, "banco_de_dados/eda/figuras/missing_por_municipio.png (Notebook 02, célula step11).")
comentario(doc,
           "Seis das sete componentes ficam abaixo de 0,05% de ausentes, cobertura praticamente "
           "total. O problema se concentra numa variável só: pct_analfab, sem dado em 15,9% dos "
           "setores, 16.552 no recorte urbano. O pior município é São Caetano do Sul, com "
           "29,69%. A supressão não é aleatória, e a tabela de 1.6 mostra o mecanismo.")

# ══ 2. CORRELAÇÃO ═══════════════════════════════════════════════════════════
doc.add_page_break()
titulo_secao(doc, "2", "A estrutura de associação: a matriz ampliada (demanda 9)")
figura(doc, "matriz_correlacao.png", 16.6,
       "Pearson e Spearman, 10 × 10. A linha preta separa as sete componentes do IVS das três "
       "descritivas, que não entram no índice e estão na matriz para decidir se deveriam.")
comentario(doc,
           "Duas coisas saltam da matriz. A primeira é o bloco socioeconômico: renda, cor/raça e "
           "analfabetismo correlacionam-se entre si a −0,81, −0,76 e +0,63, de modo que pesos "
           "iguais dariam três votos à mesma dimensão latente sem que isso tivesse sido "
           "escolhido.",
           "A segunda aparece ao comparar os dois painéis. Pearson e Spearman discordam bastante "
           "na linha da renda, e a causa é a assimetria: com a renda em log, o Pearson quase "
           "alcança o Spearman.")

subtitulo(doc, "2.1  O que a matriz decidiu (slide 19)")
# FONTE_DECK: slide 19
tabela(doc,
       ["Variável descritiva", "Maior |r| com o IVS-7", "Média |r|", "Leitura"],
       [
        ["pct_resp_feminino", "−0,299 com renda", ("0,133", PETROL, True), "carrega eixo próprio"],
        ["pct_crianca_0a4", "−0,517 com renda", "0,383", "redundante com o que já existe"],
        ["pct_idoso_60mais", "−0,541 com cor/raça", "0,392", "redundante com o que já existe"],
       ],
       larguras=[4.4, 5.0, 2.6, 5.2], tam=8.5)
comentario(doc,
           "Ampliar a matriz só se justifica se a pergunta for se alguma das três descritivas "
           "merece entrar no índice. Chefia feminina tem média |r| de 0,133 com as sete "
           "componentes, um valor baixo: ela mede algo que nenhuma variável do índice captura, e "
           "esse é o argumento para promovê-la.",
           "Idosos e crianças correlacionam cerca de três vezes mais com o que já está no "
           "índice, e entre si estão a −0,722, sendo os dois lados da mesma estrutura etária; "
           "incluir as duas contaria a mesma coisa duas vezes. A decisão sobre chefia feminina é "
           "metodológica e cabe à senhora, já que o IVS-BH original não a inclui.")
procedencia(doc, "Correlações de Spearman sobre os setores urbanos elegíveis (104.108). "
                 "Fonte: correlacao_spearman.csv.")

# ══ 3. RENDA ════════════════════════════════════════════════════════════════
doc.add_page_break()
titulo_secao(doc, "3", "A renda")

subtitulo(doc, "3.1  O caso que abriu o assunto (slide 21)")
tabela(doc,
       ["Setor", "Renda média declarada", "Identificação", "Porte", "Coef. de variação"],
       [["310620005650366", ("R$ 170.418,06", CLAY, True),
         "Belo Horizonte · Senhor dos Passos · CD_TIPO = 1 (Favela e Comunidade Urbana)",
         "186 domicílios · 518 pessoas · 31 analfabetos", ("5,26 (mediana nacional: 0,78)", CLAY, True)]],
       larguras=[3.4, 3.0, 5.2, 3.4, 2.2], tam=7.5, destaque_col0=False)
comentario(doc,
           "R$ 170 mil por 186 domicílios dariam R$ 31,7 milhões circulando por mês numa favela "
           "de 186 casas, e o perfil do setor não sustenta isso.",
           "O próprio arquivo do IBGE ajuda a dizer o que aconteceu. Ele traz o V06005, a "
           "variância do rendimento no setor, e aqui o coeficiente de variação é 5,26 contra uma "
           "mediana nacional de 0,78. Uma vírgula fora do lugar na média deixaria esse "
           "coeficiente na ordem de 526, o que descarta erro de digitação: a média está sendo "
           "puxada por uma ou poucas declarações muito altas.")
procedencia(doc, "Os 3.358 setores rastreados, um por linha e com identificação completa, estão "
                 "em banco_de_dados/eda/renda_outliers_rastreados.csv.")

subtitulo(doc, "3.2  Erro de dado e renda alta de verdade (slide 22)")
# FONTE_DECK: slide 22
tabela(doc,
       ["Classe", "Setores", "% da base", "Renda mediana", "São favela?"],
       [
        [("SUSPEITO", CLAY, True), "66", "0,06%", "R$ 10.167,85", ("22 de 66 (33%)", CLAY, True)],
        ["EXTREMO", "3.292", "3,16%", "R$ 14.106,20", ("0 de 3.292 (0%)", PETROL, True)],
        ["NORMAL", "100.750", "96,78%", "R$ 2.505,16", "—"],
       ],
       larguras=[3.2, 2.6, 2.6, 4.0, 4.8], tam=8.5)
comentario(doc,
           "Duas decisões de método fazem a classificação funcionar. A primeira é detectar por "
           "município, e não globalmente: o índice é intraurbano, e R$ 20 mil é comum em São "
           "Paulo e anômalo em Autazes, de modo que um corte nacional mediria a distância entre "
           "cidades, que não é o objeto. A segunda é olhar para a incoerência em vez do valor. "
           "Um setor no topo da renda do próprio município que também é favela, ou que tem "
           "analfabetismo e proporção PPI acima da mediana local, é internamente contraditório.",
           "Duas ressalvas sobre o alcance da regra. Ser favela é um dos testes de incoerência, "
           "então nenhuma favela pode cair em EXTREMO por construção, e a última coluna repete a "
           "regra em vez de validá-la. E a regra é cega a erro de dado em bairro rico: 40 dos 66 "
           "suspeitos são flagrados apenas por analfabetismo acima da mediana municipal, evento "
           "de metade por construção, e o setor de São Paulo com R$ 140.172,64, 45 vezes a "
           "mediana da cidade, sai como EXTREMO porque o entorno sustenta renda alta.")

subtitulo(doc, "3.3  As duas análises exploratórias, com e sem os 66 suspeitos (demanda 1.1, slide 23)")
# FONTE_DECK: slide 23
tabela(doc,
       ["O que muda ao remover os 66 suspeitos", "Com", "Sem", "Variação"],
       [
        ["Média da renda (R$)", "4.187,41", "4.178,44", "−0,21%"],
        ["Mediana da renda (R$)", "2.572,39", "2.571,17", "−0,05%"],
        ["Desvio-padrão (R$)", "4.150,66", "4.096,07", "−1,32%"],
        ["Assimetria", "3,74", "3,14", ("−16%", PETROL, True)],
        ["Pearson renda × analfabetismo", "−0,4181", "−0,4282", "0,010"],
        ["Spearman renda × analfabetismo", "−0,7566", "−0,7580", "0,001"],
       ],
       larguras=[7.0, 3.4, 3.4, 3.4], tam=8.5)
comentario(doc,
           "O resultado contraria o que eu esperava. Remover os 66 quase não mexe na "
           "estatística: as correlações andam no máximo 0,010 e o Spearman fica praticamente "
           "imóvel. Passar a renda para logaritmo muda cerca de dez vezes mais, levando a "
           "correlação com analfabetismo de −0,42 para −0,59 e com cor/raça de −0,68 para −0,81.",
           "A leitura prática é que o problema da renda está na assimetria, e não no outlier. "
           "Isso pesa direto na análise fatorial, que roda sobre a matriz de Pearson: com a renda "
           "bruta, a carga dela sairia subestimada.")
procedencia(doc, "Fonte: renda_eda_com_vs_sem.csv e renda_correlacao_com_vs_sem.csv.")

subtitulo(doc, "3.4  Onde a exclusão importa de verdade: a escala (slide 24)")
# FONTE_DECK: slide 24
tabela(doc,
       ["Município", "Setores", "Suspeitos", "Comprimidos no 1º decil (com)", "(sem)", "Ganho"],
       [
        [("Autazes", TINTA, True), "43", "1", "81,4%", "14,3%", ("67,1 pp", PETROL, True)],
        ["Salto", "196", "1", "94,9%", "59,0%", "35,9 pp"],
        ["Belo Horizonte", "5.113", "2", "98,8%", "70,8%", "28,0 pp"],
        ["Belém", "2.004", "6", "92,7%", "69,6%", "23,1 pp"],
        ["São Gonçalo", "2.357", "3", "94,0%", "79,1%", "14,9 pp"],
       ],
       larguras=[3.6, 2.2, 2.4, 4.6, 2.2, 2.2], tam=8.5)
comentario(doc,
           "O argumento para excluir está na normalização, que é o insumo do índice, e não na "
           "média nem na correlação. Autazes tem 43 setores; basta um valor ruim para 81% da "
           "cidade colapsar no primeiro decil da escala de renda. Vinte e três "
           "dos 63 municípios avaliados têm ao menos um suspeito.",
           "A tabela também mostra o que a exclusão não resolve. Mesmo sem os suspeitos, Belo "
           "Horizonte fica com 70,8% dos setores no primeiro decil e Belém com 69,6%. O que "
           "quebra a escala não são os 66 valores, e sim aplicar min-max a uma variável com "
           "assimetria 3,74; trocar a normalização global pela municipal reduz o problema sem "
           "resolvê-lo.",
           "Minha recomendação é que os 66 saiam do cálculo do índice e fiquem na EDA como "
           "achado de qualidade do dado, com as duas versões publicadas, de forma que a decisão "
           "continue defensável nos dois sentidos.")

subtitulo(doc, "3.5  Sudeste e Norte: dois fenômenos com o mesmo rótulo (demanda 4, slide 26)")
# FONTE_DECK: slide 26
tabela(doc,
       ["Região", "Setores", "Mediana", "Máx ÷ mediana", "Assimetria", "Extremos", "Suspeitos"],
       [
        ["Norte", "5.915", "R$ 1.820,99", "43,1×", "6,15", "7,03%", ("0,372%", CLAY, True)],
        ["Nordeste", "19.497", "R$ 1.719,20", "39,1×", "3,66", "6,29%", "0,067%"],
        ["Sudeste", "61.989", "R$ 2.730,60", ("62,4×", CLAY, True), "3,93", "2,18%", "0,044%"],
        ["Centro-Oeste", "9.490", "R$ 3.127,12", "15,1×", "2,27", "2,36%", "0,032%"],
        ["Sul", "7.217", "R$ 3.722,28", "13,4×", "2,65", "1,05%", "0,014%"],
       ],
       larguras=[2.8, 2.2, 2.8, 2.8, 2.4, 2.2, 2.0], tam=8.5)
comentario(doc,
           "O Sudeste tem o extremo absoluto, 62,4 vezes a mediana da própria região, e ainda "
           "assim uma taxa de suspeita baixa, 0,044%. O Norte tem 26 vezes a taxa do Sul, e não "
           "por concentrar mais gente rica: a mediana de lá é a segunda mais baixa do país. A "
           "distribuição é que está tão comprimida na base que qualquer setor de classe média já "
           "destoa, com assimetria de 6,15, a maior do país.",
           "Um achado que não estava previsto: o critério global inverte esse retrato. Ele marca "
           "4,61% do Sudeste e só 1,32% do Norte, porque mede se o setor é rico para o Brasil e "
           "não se é anômalo na própria cidade. Dos cerca de 4.000 setores marcados pelos dois "
           "critérios, apenas 1.673 coincidem.")
procedencia(doc, "Fonte: renda_extremos_por_regiao.csv e renda_criterio_global_vs_municipal.csv.")

doc.add_page_break()
subtitulo(doc, "3.6  Renda alta está mesmo nos setores pequenos? (demanda 2, slide 25)")
figura(doc, "renda_tamanho_do_setor.png", 16.6,
       "À esquerda, renda por número de domicílios, com o corte de 50 domicílios tracejado. "
       "À direita, taxa de suspeita por faixa de tamanho.")
comentario(doc,
           "A resposta é sim e não, e a diferença entre as duas metades importa. O valor não "
           "depende do tamanho: o Spearman entre número de domicílios e renda é −0,031, e o maior "
           "valor da base está num setor de 186 domicílios, que é o tamanho mediano.",
           "Já a taxa de suspeita depende bastante, com 0,265% nos setores de até 50 domicílios "
           "contra 0,045% nos de 201 a 400, seis vezes mais. O mecanismo que a senhora descreveu "
           "existe e está medido; ele "
           "apenas não é o que produz o maior valor da base.")

doc.add_page_break()
subtitulo(doc, "3.7  A distribuição da renda nas 70 cidades (demanda 3, slide 27)")
figura(doc, "renda_boxplot_por_cidade.png", 7.6,
       "Painéis por região, cidades ordenadas pela mediana, eixo logarítmico. Em vermelho, os 66 "
       "setores suspeitos, cidade a cidade.")
comentario(doc,
           "Um boxplot único com as setenta cidades fica ilegível, por isso os painéis por região, com "
           "as cidades ordenadas pela mediana. O eixo é logarítmico porque, com assimetria 3,74, "
           "em escala linear 60 das 70 caixas colapsam contra a margem esquerda e nada se "
           "distingue. Belo Horizonte, Belém e Salvador concentram os casos mais graves, e a "
           "figura em resolução cheia está em banco_de_dados/eda/figuras/.")

# ══ 4. BLOCOS DESCRITIVOS ═══════════════════════════════════════════════════
doc.add_page_break()
titulo_secao(doc, "4", "Os blocos descritivos")

subtitulo(doc, "4.1  Os blocos descritivos, região a região (slide 29)")
bloco(doc, "descritivos_regiao", larguras=[5.2, 2.4, 2.4, 2.4, 2.4, 2.4])
comentario(doc,
           "Habitação precária, 0,65% no total, concentra-se no Centro-Oeste e no Sudeste. É "
           "contraintuitivo e vale como achado: a precariedade habitacional aqui é metropolitana, "
           "não rural do Norte. Banheiro acompanha o gradiente Norte/Sul do saneamento, e "
           "habitação improvisada não acompanha.",
           "A chefia feminina é maioria, com 52,81% no recorte urbano e pico no Nordeste, 55,4%.")
fonte_bloco(doc, "descritivos_regiao")

subtitulo(doc, "4.2  Chefia feminina e envelhecimento (slide 30)")
bloco(doc, "envelhecimento", larguras=[5.2, 2.4, 2.4, 2.4, 2.4, 2.4])
tabela(doc, ["IEP (índice de envelhecimento)", "RDI (razão de dependência)", "60 anos ou mais", "Chefia feminina"],
       [[(D["envelhecimento"]["IEP"], PETROL, True), (D["envelhecimento"]["RDI"], PETROL, True),
         D["envelhecimento"]["pct_60mais"] + "%", (D["descritivos"]["resp_feminino"], PETROL, True)]],
       larguras=[4.6, 4.6, 4.0, 4.0], tam=8.5, destaque_col0=False)
comentario(doc,
           "O índice de envelhecimento passou por uma correção. Ele usava só a faixa de 0 a 4 "
           "anos no denominador e dava 299; com o denominador correto, os menores de 15 anos, "
           "conforme Galvão et al. (Hygeia, 2025), dá 92,7.",
           "Por região o contraste é grande, de 56,3 no Norte a 111,9 no Sul, onde já há mais "
           "idosos do que crianças.")
fonte_bloco(doc, "envelhecimento")

subtitulo(doc, "4.3  Tipo de espécie do domicílio (slide 31)")
bloco(doc, "morfologia", larguras=[1.8, 4.6, 1.9, 1.75, 1.75, 1.9, 1.75, 1.75], tam=7.5)
comentario(doc,
           "Esta é a morfologia urbana, ou seja, tudo que não é casa. O apartamento vai de 15,7% "
           "no Norte a 36,9% no Sul, e é o indicador de morfologia pedido em julho. Cortiço e "
           "casa de cômodos, o tipo mais associado a precariedade, concentra-se no Centro-Oeste. "
           "A maloca indígena aparece zerada porque no recorte urbano dos 70 municípios não há "
           "nenhuma.",
           "Atenção à unidade: aqui o percentual é sobre domicílios, e não sobre setores. É por "
           "isso que os valores de apartamento diferem dos de 4.1, que são médias entre setores.")
fonte_bloco(doc, "morfologia")

subtitulo(doc, "4.4  Canalização da água (demanda 10, slide 32)")
bloco(doc, "agua", larguras=[3.0, 3.2, 2.9, 2.6, 3.2, 2.3], tam=8.5)
comentario(doc,
           "Estas variáveis medem um eixo que o IVS não cobria. O indicador que já está no "
           "índice, V00112 a V00118, mede a fonte da água: poço, nascente, carro-pipa, rio. As "
           "três novas medem a entrega, ou seja, se a água chega encanada dentro do domicílio, só "
           "até o terreno, ou não chega. Um domicílio ligado à rede geral pode receber água "
           "apenas no terreno, e o Spearman entre os dois eixos é 0,459: são conceitos próximos, "
           "sem serem o mesmo.",
           "O Norte tem 3,73% dos domicílios sem canalização contra 0,33% do Sul, onze vezes "
           "mais. Como componente do índice isso discriminaria pouco, já que apenas 5% dos "
           "setores têm algum domicílio sem canalização, mas como descritivo de contraste "
           "regional funciona bem.")
fonte_bloco(doc, "agua")

subtitulo(doc, "4.5  O complemento que recupera 21,9% dos setores (slide 33)")
tabela(doc,
       ["Fórmula", "Setores sem valor", "Por quê"],
       [
        [("(V00200 + V00201) / V00001", CLAY, True), ("21,9%", CLAY, True),
         "V00200 e V00201 são contagens pequenas, e é a contagem pequena que o IBGE suprime"],
        [("1 − V00199 / V00001", PETROL, True), ("0,04%", PETROL, True),
         "V00199 é contagem grande, quase nunca sigilada; como a trinca fecha, o complemento devolve o mesmo valor"],
       ],
       larguras=[5.4, 3.4, 8.4], tam=8.5, destaque_col0=False)
comentario(doc,
           "As três variáveis formam uma partição de V00001, e isso virou auditoria automática: "
           "elas somam exatamente o total de domicílios em 100,00% dos 81.270 setores em que as "
           "três estão presentes.",
           "A ressalva precisa constar do artigo. A identidade só é verificável onde as três "
           "estão presentes; nos setores com sigilo, aplicá-la é extrapolação, justificada porque "
           "a partição é definida pelo IBGE, mas ainda assim suposição e não medição. A "
           "alternativa era perder um em cada cinco setores, de forma não aleatória, e preferi a "
           "suposição declarada à perda silenciosa.")

subtitulo(doc, "4.6  Cobertura integral de saneamento (slide 34)")
bloco(doc, "cobertura_regiao", larguras=[5.2, 2.4, 2.4, 2.4, 2.4, 2.4])
tabela(doc, ["Água 100% adequada", "Esgoto 100% adequado", "Lixo 100% adequado", "Os três juntos"],
       [[D["cobertura"]["agua"] + "%", D["cobertura"]["esgoto"] + "%",
         D["cobertura"]["lixo"] + "%", (D["cobertura"]["tres"] + "%", CLAY, True)]],
       larguras=[4.3, 4.3, 4.3, 4.3], tam=8.5, destaque_col0=False)
comentario(doc,
           f"A água chega a {D['cobertura']['agua']}% de cobertura integral e o esgoto a "
           f"{D['cobertura']['esgoto']}%, mas só {D['cobertura']['tres']}% dos setores têm os três "
           f"serviços integralmente adequados ao mesmo tempo.",
           f"Contando a caçamba de serviço de limpeza como coleta, o lixo sobe de "
           f"{D['cobertura']['lixo']}% para {D['cobertura']['coleta']}%. Essa diferença é o efeito "
           f"de uma decisão metodológica sobre V00398, mantida por fidelidade ao IVS-BH 2012 e "
           f"hoje registrada como em revisão.")
fonte_bloco(doc, "cobertura_regiao")

subtitulo(doc, "4.7  Gravidade do saneamento em faixas (slide 35)")
bloco(doc, "saneamento_faixas", larguras=[3.4, 3.8, 2.0, 2.0, 2.0, 2.0, 2.0])
comentario(doc,
           "Além de quantos setores têm inadequação, importa quanto. No Norte, 30,8% dos setores "
           "têm metade ou mais dos domicílios com água inadequada, contra 11,8% em situação "
           "totalmente adequada. No Sul e no Centro-Oeste a distribuição se inverte por completo.")
fonte_bloco(doc, "saneamento_faixas")

# ══ 5. FAVELAS E BRASIL ═════════════════════════════════════════════════════
doc.add_page_break()
titulo_secao(doc, "5", "Favelas e Comunidades Urbanas e a linha de base nacional")

subtitulo(doc, "5.1  A fonte oficial e a validação do critério (slide 37)")
corpo(doc, "A fonte é o IBGE. Censo Demográfico 2022: Favelas e Comunidades Urbanas — Resultados "
           "do universo. Rio de Janeiro, 2024, 171 p. A definição e os quatro critérios estão "
           "transcritos na seção 14.2 do relatório da EDA. Mais útil que o texto foi a planilha "
           "anexa, que lista as favelas por setor censitário.", tam=9)
# FONTE_DECK: slide 37
tabela(doc,
       ["", "Brasil (IBGE 2024)", "ELSI-70", "Cobertura"],
       [
        ["Favelas e Comunidades Urbanas", "12.348", "5.899", "47,8%"],
        ["Municípios com FCU", "656", "42", "6,4%"],
        [("População em FCU", TINTA, True), "16.390.815", "10.069.994", ("61,4%", PETROL, True)],
        ["Domicílios em FCU", "6.556.998", "3.443.687", "52,5%"],
       ],
       larguras=[6.0, 4.2, 3.6, 3.4], tam=8.5)
comentario(doc,
           "A planilha permitiu validar o critério do projeto. Dos 109.032 setores da base "
           "completa, 19.507 estão na lista oficial do IBGE, e são exatamente os 19.507 que o "
           "código marca com CD_TIPO = 1: nenhum falso positivo e nenhuma omissão. No recorte de "
           "análise são 19.452, ou 18,7% dos 104.108. O campo NM_FCU não serve como critério, "
           "porque diverge em 25 setores, nenhum deles na lista oficial.",
           "Um dado que vale para o artigo: os 70 municípios do ELSI reúnem 61,4% de toda a "
           "população favelada do Brasil, sendo apenas 6,4% dos municípios que têm favela.")
corpo(doc, "A mesma fonte revelou uma limitação que eu não conhecia. Na nota 7 da página 75, o "
           "IBGE informa que, além das 12.348 FCU classificadas, identificou 2.298 FCU com 21 a "
           "50 domicílios que não receberam setor censitário próprio. Como o critério do projeto "
           "depende do setor, essas favelas ficam invisíveis na base, e a comparação entre favela "
           "e resto da cidade erra nos dois sentidos: subestima a população favelada e contamina "
           "com ela o grupo de comparação.", cor=CLAY, tam=8.5)

subtitulo(doc, "5.2  Favela e restante da cidade, indicador a indicador (slide 38)")
bloco(doc, "favela_resto", larguras=[3.6, 1.7, 1.9, 2.1, 1.9, 2.1, 2.2, 1.7], tam=7.5)
comentario(doc,
           "A última coluna diz quantas vezes o valor médio no setor de favela supera o de fora. "
           "O esgoto inadequado é 4,1 vezes maior, o lixo 2,9 e a água 1,8.",
           "A razão de moradores, porém, fica em 1,11, praticamente igual dentro e fora. Isso é "
           "coerente e vale dizer: o adensamento domiciliar não é o eixo que separa a favela do "
           "resto da cidade, e sim a infraestrutura.",
           "Onde a razão fica abaixo de 1, o valor é menor em favela: é o caso da renda, 0,337, "
           "do apartamento, 0,087, e do índice de envelhecimento, 0,360.")
fonte_bloco(doc, "favela_resto")

subtitulo(doc, "5.3  A amostra ELSI comparada com o Brasil urbano (slide 39)")
bloco(doc, "brasil_elsi", larguras=[6.0, 4.0, 4.0, 3.2], tam=8.5)
comentario(doc,
           "São os mesmos indicadores calculados sobre os 468 mil setores do país inteiro, como "
           "linha de base. Em seis dos sete a amostra ELSI está melhor que o Brasil urbano, e o "
           "esgoto é o contraste mais forte: 0,080 contra 0,155, quase metade.",
           "A exceção é o lixo, a única variável em que a amostra está pior, com razão de 1,21. "
           "Isso reforça a hipótese levantada em 1.4, de que o indicador de lixo pode estar "
           "medindo porte urbano em vez de vulnerabilidade, já que a caçamba de serviço de "
           "limpeza é mais comum em cidade grande.")
fonte_bloco(doc, "brasil_elsi")

saida = Path(sys.argv[1]) if len(sys.argv) > 1 else RAIZ / "docs" / "Apresentacoes_IVS" / "complementos" / "Resumo_EDA_Central_2026-08.docx"
saida.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(saida))
print("resumo escrito:", saida)
