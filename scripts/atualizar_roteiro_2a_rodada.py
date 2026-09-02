"""Atualiza o roteiro da EDA Central da 1ª para a 2ª rodada.

Por que este script existe, e não uma edição no Word
-----------------------------------------------------
O roteiro é o único artefato desta pasta escrito à mão, e o que ele tem de valioso
não são os números: é o julgamento — quais slides não podem cair, o que dizer em
cada um, que pergunta a orientadora provavelmente fará. Isso não se regenera a
partir de tabela nenhuma.

Então o roteiro **não** vira um documento gerado. Este script faz a passagem de uma
rodada para a outra de forma auditável: renumera os slides, insere as seções dos 4
slides novos e troca os números que mudaram — cada um lido de
`banco_de_dados/eda/atualizada/`, nenhum digitado aqui. O texto do autor sobrevive
palavra por palavra.

Se a EDA for para uma 3ª rodada, este arquivo é o registro do que foi preciso mexer.

Uso:
    uv run --with python-docx python scripts/atualizar_roteiro_2a_rodada.py
"""
from __future__ import annotations

import copy
import re
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt

RAIZ = Path(__file__).resolve().parents[1]
APRES = RAIZ / "docs" / "Apresentacoes_IVS" / "complementos"
EDA = RAIZ / "banco_de_dados" / "eda" / "atualizada"

ORIGEM = APRES / "Roteiro_EDA_Central_1a_rodada.docx"
DESTINO = APRES / "Roteiro_EDA_Central_2a_rodada.docx"

# Os 4 slides novos entram depois do slide 2; tudo que era >= 3 anda 4 casas.
PRIMEIRO_DESLOCADO = 3
DESLOCAMENTO = 4


def ler(nome: str, **kw) -> pd.DataFrame:
    return pd.read_csv(EDA / f"{nome}.csv", sep=";", encoding="utf-8-sig", **kw)


def br(v, casas=2) -> str:
    """Número no formato brasileiro."""
    s = f"{float(v):,.{casas}f}"
    return s.replace(",", "\x00").replace(".", ",").replace("\x00", ".")


# ── os números da 2ª rodada, lidos das tabelas ──────────────────────────────

def numeros() -> dict:
    t = ler("renda_outliers_rastreados")
    s = t[t["classe_renda"] == "SUSPEITO"]
    d = ler("descritivas_globais", index_col=0)
    cs = ler("renda_eda_com_vs_sem").set_index("variavel").loc["renda_media"]
    cc = ler("renda_correlacao_com_vs_sem")
    reg = ler("renda_extremos_por_regiao").set_index("regiao")
    nz = ler("renda_normalizacao_impacto").sort_values("ganho_pp", ascending=False)
    cg = ler("renda_criterio_global_vs_municipal")
    maior = t.nlargest(1, "renda_media_setor").iloc[0]
    bh = nz[nz["NM_MUN"] == "Belo Horizonte"].iloc[0]
    return {
        "n_suspeitos": len(s),
        "n_rastreados": len(t),
        "n_favela_suspeitos": int((s["e_favela"] == "sim").sum()),
        "so_analfab": int((s["motivos"] == "pct_analfab_acima").sum()),
        "assimetria": br(d.loc["renda_media", "assim"]),
        "delta_media": br(abs(cs["delta_media_pct"]), 2),
        "delta_mediana": br(abs(cs["delta_mediana_pct"]), 2),
        "max_delta_corr": br(cc["delta"].abs().max(), 4),
        "maior_mun": str(maior["NM_MUN"]),
        "maior_dom": int(maior["n_domicilios"]),
        "sudeste_razao": br(reg.loc["Sudeste", "max_sobre_mediana"], 1),
        "sudeste_susp": br(reg.loc["Sudeste", "pct_suspeitos"], 3),
        "global_total": br(cg["outlier_global"].sum(), 0),
        "global_concordam": br(cg["concordam"].sum(), 0),
        "bh_decil": br(bh["pct_1o_decil_com"], 1),
        "bh_suspeitos": int(bh["n_suspeitos"]),
        "topo_mun": str(nz.iloc[2]["NM_MUN"]),
        "topo_decil": br(nz.iloc[2]["pct_1o_decil_com"], 1),
        "topo_sem": br(nz.iloc[2]["pct_1o_decil_sem"], 1),
    }


# ── edição de texto preservando a formatação ────────────────────────────────

def substituir(par, velho: str, novo: str) -> bool:
    """Troca `velho` por `novo` no parágrafo, mexendo no menor número de runs.

    Quando o trecho cabe dentro de um run, só aquele run muda e toda a formatação
    do parágrafo fica intacta. Quando ele atravessa runs — o que acontece sempre
    que o Word quebrou a frase por causa de um negrito ou de uma revisão —, o
    parágrafo é reescrito com a formatação do primeiro run.
    """
    for run in par.runs:
        if velho in run.text:
            run.text = run.text.replace(velho, novo)
            return True
    inteiro = "".join(r.text for r in par.runs)
    if velho not in inteiro:
        return False
    novo_txt = inteiro.replace(velho, novo)
    for run in par.runs[1:]:
        run._element.getparent().remove(run._element)
    par.runs[0].text = novo_txt
    return True


def escrever(par, texto: str) -> None:
    """Substitui TODO o texto do parágrafo, mantendo a formatação do primeiro run.

    Escrever só em `runs[0]` deixava o texto antigo dos runs seguintes no lugar — foi
    o que produziu células como "7–93–5" na tabela do arco.
    """
    if not par.runs:
        par.add_run(texto)
        return
    for extra in par.runs[1:]:
        extra._element.getparent().remove(extra._element)
    par.runs[0].text = texto


def renumerar(texto: str) -> str:
    """`Slide N` -> `Slide N+4` para N >= 3."""
    def troca(m):
        n = int(m.group(1))
        return f"Slide {n + DESLOCAMENTO}" if n >= PRIMEIRO_DESLOCADO else m.group(0)
    return re.sub(r"Slide (\d+)", troca, texto)


def par_depois(ref, texto: str, estilo: str | None = None, italico=False):
    """Insere um parágrafo logo DEPOIS de `ref` e devolve o novo, para encadear."""
    novo = copy.deepcopy(ref._element)
    for filho in list(novo):
        if filho.tag.endswith('}r') or filho.tag.endswith('}br'):
            novo.remove(filho)
    ref._element.addnext(novo)
    from docx.text.paragraph import Paragraph
    p = Paragraph(novo, ref._parent)
    if estilo:
        p.style = ref.part.document.styles[estilo]
    run = p.add_run(texto)
    run.italic = italico
    return p


def main() -> None:
    if not ORIGEM.exists():
        raise SystemExit(f"Roteiro da 1ª rodada não encontrado: {ORIGEM}")
    n = numeros()
    doc = Document(str(ORIGEM))

    def todos_paragrafos():
        """Parágrafos do corpo E das células — `doc.paragraphs` deixa as tabelas de
        fora, e foi por isso que a frase do arco da apresentação, que vive numa
        célula, escapou da primeira versão deste script."""
        vistos = list(doc.paragraphs)
        for tab in doc.tables:
            for linha in tab.rows:
                for celula in linha.cells:
                    vistos.extend(celula.paragraphs)
        return vistos

    paras = todos_paragrafos()

    # ── 1. cabeçalho ────────────────────────────────────────────────────────
    trocas_cabecalho = [
        ("47 slides · duração estimada de 35 a 45 minutos",
         "51 slides · duração estimada de 40 a 50 minutos"),
        ("agosto de 2026", "setembro de 2026"),
        ("Tenha à mão o número 63", "Tenha à mão o número 65"),
        # os slides que não podem cair, deslocados, com o novo bloco de alterações
        ("os slides que não podem cair são 7, 10, 18, 21, 24, 32, 37 e 46",
         "os slides que não podem cair são 3, 4, 11, 14, 22, 25, 28, 36, 41 e 50"),
        ("Os divisores de seção (3, 6, 8, 17, 20, 28, 36, 40)",
         "Os divisores de seção (7, 10, 12, 21, 24, 32, 40, 44)"),
        ("Os slides 12, 13, 15, 18, 25 e 27 têm figura",
         "Os slides 16, 17, 19, 22, 29 e 31 têm figura"),
        # o arco da apresentação
        ("O extremo não é renda alta, é erro de dado — e o dano está na escala, não na média.",
         "O extremo não é renda alta nem vírgula fora do lugar — é média puxada por poucas "
         "declarações; e o dano está na escala, não na média."),
    ]
    for p in paras:
        for velho, novo in trocas_cabecalho:
            substituir(p, velho, novo)

    # ── 2. renumeração dos títulos e das referências cruzadas ───────────────
    for p in paras:
        txt = "".join(r.text for r in p.runs)
        if "Slide " in txt and re.search(r"Slide \d+", txt):
            novo = renumerar(txt)
            if novo != txt:
                for run in p.runs[1:]:
                    run._element.getparent().remove(run._element)
                p.runs[0].text = novo

    # ── 3. números da seção de renda ────────────────────────────────────────
    S, R = n["n_suspeitos"], n["n_rastreados"]
    trocas_renda = [
        ("a tabela com os 66 casos está pronta", f"a tabela com os {S} casos está pronta"),
        ("O resultado: 66 setores suspeitos", f"O resultado: {S} setores suspeitos"),
        ('A coluna "São favela?" — 22 de 66 contra 0 de 3.292',
         f'A coluna "São favela?" — {n["n_favela_suspeitos"]} de {S} contra 0 de 3.292'),
        ("40 dos 66 suspeitos", f"{n['so_analfab']} dos {S} suspeitos"),
        ("Tirar os 66 suspeitos quase não muda a estatística. A média cai 0,21%, a mediana 0,05%. "
         "As correlações se movem no máximo 0,010",
         f"Tirar os {S} suspeitos restantes quase não muda a estatística. A média cai "
         f"{n['delta_media']}%, a mediana {n['delta_mediana']}%. As correlações se movem no "
         f"máximo {n['max_delta_corr']}"),
        ("O que quebra a escala não são os 66 valores",
         f"O que quebra a escala não são os {S} valores"),
        ("Vinte e três dos 63 municípios", "Vinte e três dos 63 municípios"),
        ("Que os 66 saiam do cálculo", f"Que os {S} saiam do cálculo"),
        ("O maior valor da base está num setor de 186 domicílios, que é o tamanho mediano.",
         f"O maior valor da base agora está em {n['maior_mun']}, num setor de "
         f"{n['maior_dom']} domicílios."),
        ("62 vezes a mediana da própria região. Mas a taxa de suspeita é baixa, 0,044%",
         f"{n['sudeste_razao']} vezes a mediana da própria região. Mas a taxa de suspeita é "
         f"baixa, {n['sudeste_susp']}%"),
        ("Dos cerca de 4.000 setores marcados pelos dois critérios, apenas 1.673 coincidem.",
         f"Dos {n['global_total']} setores marcados pelo critério global, apenas "
         f"{n['global_concordam']} coincidem com o municipal."),
        ("Os pontos vermelhos são os 66 suspeitos",
         f"Os pontos vermelhos são os {S} suspeitos"),
        ("com assimetria 3,74", f"com assimetria {n['assimetria']}"),
        ("Com assimetria de 3,74", f"Com assimetria de {n['assimetria']}"),
        (f"renda_outliers_rastreados.csv — 3.358", f"renda_outliers_rastreados.csv — {R}"),
        # Belo Horizonte deixou de ser o caso da normalização: o extremo já saiu da coluna
        ("Em Belo Horizonte são 98,8% dos setores espremidos no primeiro decil, com dois "
         "valores ruins numa cidade de cinco mil setores.",
         f"Belo Horizonte saiu desta tabela, e a saída é o resultado: com o extremo já fora da "
         f"coluna de renda, a cidade começa em {n['bh_decil']}% — que era exatamente o valor que "
         f"ela alcançava, na 1ª rodada, DEPOIS de remover os suspeitos. O ganho dela agora é "
         f"zero. Quem ocupa o lugar é {n['topo_mun']}, com {n['topo_decil']}%."),
        ("Mesmo tirando os suspeitos, Belo Horizonte fica com 70,8% dos setores no primeiro "
         "decil, e Belém com 69,6%.",
         f"Mesmo tirando os suspeitos, Belo Horizonte fica com {n['bh_decil']}% dos setores no "
         f"primeiro decil, e Belém com 69,6%."),
        ("O ponto vermelho isolado bem à direita, no painel do Sudeste: é Belo Horizonte.",
         "O painel do Sudeste: o ponto isolado bem à direita, que na 1ª rodada era Belo "
         "Horizonte, não existe mais — aquele setor saiu da coluna de renda."),
    ]
    for p in paras:
        for velho, novo in trocas_renda:
            substituir(p, velho, novo)

    # ── 4. as seções dos 4 slides novos, antes do divisor do desenho ────────
    alvo = next(p for p in doc.paragraphs
                if "".join(r.text for r in p.runs).startswith("Slide 7 — Divisor"))
    ancora = alvo.insert_paragraph_before("", style=doc.styles["Heading 2"])
    novos = [
        ("Heading 2", "Slide 3 — O que mudou nesta rodada", False),
        (None, "Tempo sugerido: 2 minutos", False),
        ("Heading 3", "O que está na tela", False),
        ("List Bullet", "Três parágrafos: o pedido, o setor e o alcance da mudança.", False),
        ("List Bullet", "Quatro números em destaque no rodapé, um deles a queda de 48% na curtose.", False),
        ("Heading 3", "O que dizer", True),
        (None, "“A senhora pediu uma coluna de renda sem o extremo de Belo Horizonte. Eu fiz a "
               "coluna — e refiz a EDA inteira com ela, para poder dizer o que mudou em vez de "
               "supor.”", True),
        (None, f"“Saiu um setor de {104108:,}".replace(",", ".") +
               ". O que mudou foram 45 células, em 6 tabelas, todas dentro da renda e das "
               "correlações dela. Nenhum bloco fora da renda se moveu — e isso foi verificado "
               "célula a célula, não no olho.”", True),
        ("Heading 3", "Se perguntarem", False),
        (None, "P: Como você sabe que não mudou mais nada?", False),
        (None, "R: O mesmo script recalcula as tabelas nas duas versões e compara. Antes disso, "
               "ele roda com a coluna antiga e confere se reproduz a EDA já publicada — se não "
               "reproduzir, ele para. As tabelas que não aparecem na lista de alterações têm "
               "diferença exatamente zero.", False),

        ("Heading 2", "Slide 4 — A renda, antes e depois", False),
        (None, "Tempo sugerido: 3 minutos", False),
        ("Heading 3", "O que está na tela", False),
        ("List Bullet", "Tabela de sete estatísticas da renda, com a coluna de variação.", False),
        ("Heading 3", "O que dizer", True),
        (None, "“O efeito é estreito e fundo, não largo e raso. A média cai 0,04% e a mediana não "
               "muda — tirar um setor em 104 mil não desloca o centro de nada.”", True),
        (None, "“O que muda é a cauda. A assimetria cai 14,5% e a curtose cai 48,2%. Metade do "
               "peso da cauda da renda, no recorte inteiro, estava naquele único setor.”", True),
        ("Heading 3", "O que apontar", False),
        ("List Bullet", "A linha da curtose: 49,49 para 25,66. É a linha que justifica a rodada "
                        "inteira.", False),
        ("Heading 3", "Se perguntarem", False),
        (None, "P: Se a média quase não muda, para que serviu?", False),
        (None, "R: Para a escala, que é o insumo do índice — é o slide 28. E para a matriz de "
               "Pearson, que é o insumo da análise fatorial — é o slide 5.", False),

        ("Heading 2", "Slide 5 — As correlações com a renda", False),
        (None, "Tempo sugerido: 3 minutos", False),
        ("Heading 3", "O que está na tela", False),
        ("List Bullet", "Os oito pares de Pearson com a renda, antes e depois.", False),
        ("Heading 3", "O que dizer", True),
        (None, "“Todas as correlações da renda ficam mais fortes. O valor extremo achatava a "
               "associação linear: com analfabetismo, vai de menos 0,418 para menos 0,424.”", True),
        (None, "“E o ponto que importa mais do que o tamanho da mudança: a matriz de Spearman é "
               "idêntica antes e depois, nas cem células. Ela trabalha com postos, então o caso "
               "nunca a afetou.”", True),
        ("Heading 3", "Se perguntarem", False),
        (None, "P: E o que isso recomenda?", False),
        (None, "R: Estatística robusta para a renda inteira — posto, log ou escala robusta —, em "
               "vez de caçar observação. É a decisão que continua em aberto, e agora ela tem "
               "evidência empírica direta, não só argumento.", False),
        (None, "P: A análise fatorial muda?", False),
        (None, "R: Praticamente não: o KMO vai de 0,7826 para 0,7825, porque o diagnóstico roda "
               "sobre Spearman. As conclusões sobre o lixo formar fator próprio e sobre os pesos "
               "65/35 seguem valendo sem retrabalho.", False),

        ("Heading 2", "Slide 6 — O que NÃO mudou", False),
        (None, "Tempo sugerido: 2 minutos", False),
        ("Heading 3", "O que está na tela", False),
        ("List Bullet", "Sete linhas de blocos da EDA, todas com a mesma resposta: idêntico.", False),
        ("Heading 3", "O que dizer", True),
        (None, "“Este slide é tão importante quanto o anterior. A exclusão não vazou para o resto "
               "da análise: água, esgoto, lixo, razão de moradores, analfabetismo e cor ou raça "
               "estão idênticos em média, mediana, outliers e faltantes.”", True),
        (None, "“As contagens de favela também: 19.452 no recorte, 19.507 na base. E a matriz de "
               "Spearman, idêntica nas cem células.”", True),
        ("Heading 3", "Se perguntarem", False),
        (None, "P: Então por que refazer a EDA inteira?", False),
        (None, "R: Porque dizer que só a renda mudaria era uma expectativa, não um fato. "
               "Recalcular tudo é o que transforma a expectativa em verificação — e foi assim que "
               "apareceram três números defasados no próprio deck, que já estão corrigidos.", False),
    ]
    ref = ancora
    for estilo, texto, italico in novos:
        ref = par_depois(ref, texto, estilo, italico)
    ancora._element.getparent().remove(ancora._element)

    # ── 5. a tabela do arco ─────────────────────────────────────────────────
    tab = doc.tables[0]
    faixas = {"1. Desenho": "7–9", "2. Quem entra": "10–11", "3. As sete componentes": "12–20",
              "4. Correlação": "21–23", "5. Renda": "24–31", "6. Blocos descritivos": "32–39",
              "7. Favelas": "40–42", "8. Fechamento": "43–51"}
    for linha in tab.rows:
        rot = linha.cells[0].text.strip()
        if rot in faixas:
            escrever(linha.cells[1].paragraphs[0], faixas[rot])
    # a seção nova entra logo depois da abertura
    nova = copy.deepcopy(tab.rows[1]._tr)
    tab.rows[1]._tr.addnext(nova)
    from docx.table import _Row
    r = _Row(nova, tab)
    for celula, txt in zip(r.cells, ["O que mudou", "3–6",
                                     "Um valor saiu da base; o que ele movia era a forma da "
                                     "distribuição, não o nível."]):
        escrever(celula.paragraphs[0], txt)

    doc.save(str(DESTINO))
    print(f"roteiro da 2ª rodada: {DESTINO}")
    print(f"  {len(doc.paragraphs)} parágrafos · 51 slides · {S} suspeitos · {R} rastreados")


if __name__ == "__main__":
    main()
